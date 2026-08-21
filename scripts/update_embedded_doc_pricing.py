from pathlib import Path

from docx import Document


DOCX_PATH = Path(__file__).resolve().parents[1] / "docs" / "Mastline_Product_Overview.docx"


def set_cell(cell, value: str) -> None:
    cell.text = value
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = 0


document = Document(DOCX_PATH)

old_pricing = (
    "Pricing should be tested against the economic unit that users feel most directly: active photographer, "
    "team seat, or monthly managed asset volume. A low-friction individual plan can establish trust and create "
    "a natural path to higher-value team, intelligence, and rights capabilities."
)
new_pricing = (
    "Mastline has four operating levels. Solo is $49 per month billed annually ($588 per year) or $59 billed "
    "monthly. Pro is $99 per month billed annually ($1,188 per year) or $119 billed monthly. Studio is $279 "
    "per month billed annually ($3,348 per year) or $339 billed monthly. Agency pricing is custom. Paid plans "
    "include the optional Mastline Sales Engine; the photographer receives 70% and Mastline receives 30% only "
    "on licenses generated inside Mastline."
)

for paragraph in document.paragraphs:
    if paragraph.text.strip() == old_pricing:
        paragraph.text = new_pricing
        break
else:
    raise RuntimeError("Pricing paragraph was not found")

pricing_table = document.tables[8]
while len(pricing_table.rows) < 7:
    pricing_table.add_row()

rows = [
    ("Offer", "Pricing", "Included"),
    ("Solo", "$49/month billed annually ($588/year) or $59 monthly", "1 photographer; shoots, assets, submissions, contacts, invoices, payments, revenue reporting, and 250 GB"),
    ("Pro", "$99/month billed annually ($1,188/year) or $119 monthly", "Everything in Solo; news monitoring; archive matching; rights monitoring; advanced revenue analytics; 1 TB"),
    ("Studio", "$279/month billed annually ($3,348/year) or $339 monthly", "Everything in Pro; up to 5 people; dispatch/review queues; roles; approvals; allocation; 5 TB shared"),
    ("Agency", "Custom", "Custom team structure, migration, API/integrations, permissions, priority support, and flexible storage"),
    ("Mastline Sales Engine", "Optional on paid plans; 70% photographer / 30% Mastline", "Share applies only to licenses generated inside Mastline"),
    ("Possible extensions", "Separately disclosed", "Storage overages, rights-recovery services, and premium automation"),
]

for row, values in zip(pricing_table.rows, rows):
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)

document.save(DOCX_PATH)
